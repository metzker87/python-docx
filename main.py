from docx import Document

from docx.enum.style import WD_STYLE_TYPE
#from docx.enum.style import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches

data = {
    "Title": "Document Title", 
    "Content": [
        {
            "Chapter": "Chapter 1 text",
            "Text": "This is my text 1 \n This is still my text",
            "Image": "pmmg.png",
            "Table": ["A", "B", "C"]
        },
        {
            "Chapter": "Chapter 1 text",
            "Text": "This is my text 2 \n This is still my text 2",
            "Image": "pmmg.png",
            "Table": ["1", "2", "3"]
        },
        {
            "Chapter": "Chapter 1 text",
            "Text": "This is my text 3 \n This is still my text 3",
            "Image": "pmmg.png",
            "Table": ["1", "2", "3"]
        },
    ],
}

document = Document()

style = document.styles
# Style Paragraph
p = style.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
p.font.name = "Calibri"
p.font.size = Pt(11)
# Style Heading 2
h2 = style.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
h2.base_style = style["Heading 2"]
h2.font.name = "Calibri"
h2.font.size = Pt(13)
h2.font.color.rgb = RGBColor(79, 129, 189)
h2.font.bold = False 
# Style Heading 3
h3 = style.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
h3.base_style = style["Heading 3"]
h3.font.name = "Calibri"
h3.font.size = Pt(12)
h3.font.color.rgb = RGBColor(79, 129, 189)
h3.font.bold = False 

document.add_heading(data.get("Title"), 0)

for content in data.get("Content"):
    # Add paragraph with h2
    document.add_paragraph(content.get("Chapter"), style="H2")
    # Add paragraph
    document.add_paragraph(content.get("Text"), style="Paragraph")
    # Add paragraph with h3
    document.add_paragraph("Image", style="H3")
    document.add_picture(content.get("Image"), width=Inches(1.25))
    """
    # Add paragraph with h3
    document.add_paragraph("Table", style="H3")
    table = document.add_table(rows=1, cols=3, style="Table Grid")
    table.aligment = WD_TAB_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Col1"
    hdr_cells[1].text = "Col2"
    hdr_cells[2].text = "Col3"
    row_cells = table.add_row().cells

    for element in enumerate(content.get("Table")):
        row_cells[index].text = element
    """

document.save("NewDocx.docx")